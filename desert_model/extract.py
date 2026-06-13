"""Extraction utilities for the source Word documents."""

from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional
from xml.etree import ElementTree as ET


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
V_NS = "{urn:schemas-microsoft-com:vml}"


def _shape_int(style: str, key: str) -> Optional[int]:
    match = re.search(key + r":(-?\d+)", style)
    if not match:
        return None
    return int(match.group(1))


def extract_docx_text_tables(path: Path) -> Dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX extraction; use the bundled Python runtime") from exc

    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables: List[Dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, 1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip().replace("\n", " | ") for cell in row.cells])
        tables.append(
            {
                "index": table_index,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "values": rows,
            }
        )
    return {
        "file": str(path),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def extract_vml_summary(path: Path) -> Dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        embeddings = [name for name in archive.namelist() if name.startswith("word/embeddings/")]
    root = ET.fromstring(xml)
    shapes: List[Dict[str, Any]] = []
    lines: List[Dict[str, Any]] = []
    for shape in root.iter(V_NS + "shape"):
        text = "".join(t.text or "" for t in shape.iter(W_NS + "t")).strip()
        style = shape.attrib.get("style", "")
        item = {
            "id": shape.attrib.get("id", ""),
            "text": text,
            "left": _shape_int(style, "left"),
            "top": _shape_int(style, "top"),
            "width": _shape_int(style, "width"),
            "height": _shape_int(style, "height"),
        }
        if text or item["left"] is not None:
            shapes.append(item)
    for index, line in enumerate(root.iter(V_NS + "line"), 1):
        lines.append(
            {
                "index": index,
                "id": line.attrib.get("id", ""),
                "from": line.attrib.get("from", ""),
                "to": line.attrib.get("to", ""),
                "style": line.attrib.get("style", ""),
            }
        )
    return {
        "file": str(path),
        "media": media,
        "embeddings": embeddings,
        "shape_count": len(shapes),
        "line_count": len(lines),
        "shapes_with_text": [item for item in shapes if item["text"]],
        "lines": lines,
        "review_notes": [
            "Maps and formulas are stored as VML/WMF/OLE objects.",
            "Use this summary to create a manually reviewed adjacency table before final solve.",
        ],
    }


def extract_sources(problem_docx: Path, attachment_docx: Path, output_dir: Path) -> Dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    outputs = {
        "problem_summary": output_dir / "problem_summary.json",
        "attachment_summary": output_dir / "attachment_summary.json",
        "attachment_vml_summary": output_dir / "attachment_vml_summary.json",
        "problem_vml_summary": output_dir / "problem_vml_summary.json",
    }
    payloads = {
        outputs["problem_summary"]: extract_docx_text_tables(problem_docx),
        outputs["attachment_summary"]: extract_docx_text_tables(attachment_docx),
        outputs["attachment_vml_summary"]: extract_vml_summary(attachment_docx),
        outputs["problem_vml_summary"]: extract_vml_summary(problem_docx),
    }
    for path, payload in payloads.items():
        with path.open("w", encoding="utf-8") as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2)
            handle.write("\n")
    return outputs
